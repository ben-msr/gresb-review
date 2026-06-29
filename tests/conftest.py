import io
import docx


def make_docx():
    """Build a structurally-faithful docx with FAKE values for tests."""
    d = docx.Document()

    def header(text, level):
        # python-docx default styles 'Heading 1/2/3' stand in for the
        # real doc's 'Header 1/2/3'; the parser matches on a level map.
        d.add_heading(text, level=level)

    def table(rows):
        t = d.add_table(rows=len(rows), cols=len(rows[0]))
        for i, row in enumerate(rows):
            for j, val in enumerate(row):
                t.cell(i, j).text = str(val)

    header("Performance: Reporting Characteristics", 1)
    header("R1 Standing Investments During Reporting Period", 2)
    table([
        ["Property Type | Country", "Number of Assets", "Floor Area (ft²)", "% of GAV"],
        ["All Use Types, All Countries", "5", "1000", "100"],
        ["Hotel | United States", "1", "600", "60"],
        ["Office: Corporate: Low-Rise Office | United States", "4", "400", "40"],
    ])

    header("Performance: Risk Assessment", 1)
    header("RA2 Technical Building Assessments", 2)
    table([
        ["Topic", "Number of Assets", "% Portfolio Coverage"],
        ["Energy", "3", "55.0"],
        ["Water", "3", "55.0"],
        ["Waste", "2", "50.0"],
    ])

    header("Performance: Energy", 1)
    header("EN1 Energy Consumption Data", 2)
    header("Hotel | United States Energy Consumption", 3)
    table([
        ["Metric", "Absolute", "Absolute", "Like-for-Like"],
        ["Metric", "Prior Year Usage (MWh)", "Reporting Year Usage (MWh)", "Prior Year Usage (MWh)"],
        ["Whole Site: Indirect Fuel", "990.48", "1046.8", "990.48"],
        ["Whole Site: Indirect Electric", "799.09", "954.86", "799.09"],
    ])

    # EN1 renewables: single-header table (was mis-read by 2-header matrix
    # logic, dropping the first data row). Stays under the EN1 H2 / Hotel H3.
    table([
        ["", "Prior Year (MWh)", "Reporting Year (MWh)"],
        ["Generated On-site and Consumed by Landlord", "0", "0"],
        ["Generated On-site and Exported by Landlord", "0", "0"],
        ["Generated On-site by Third Party or Tenant", "12.5", "8.0"],
        ["Generated Off-site and Purchased by Landlord", "0", "0"],
        ["Generated Off-site and Purchased by Tenant", "0", "0"],
    ])

    header("Performance: Building Certifications", 1)
    header("BC2 Ratings by Property Type", 2)
    header("Hotel | United States", 3)
    table([
        ["Energy Rating", "ft² Certified", "% Portfolio Covered in Reporting Year", "Number of Assets"],
        ["Energy Star Portfolio Manager", "600", "100", "1"],
    ])

    buf = io.BytesIO()
    d.save(buf)
    return buf.getvalue()
